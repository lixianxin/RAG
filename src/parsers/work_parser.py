from pathlib import Path
from token import OP
from typing import List, Dict, Optional

import docx  # python-docx

from src.parsers.base import BaseParser, ParseResult
from src.parsers.chunker import TextChunker
from src.parsers.cleaner import TextCleaner
from src.parsers.table_parser import TableProcessor


class WorkParser(BaseParser):
    def __init__(
        self,
        chunker: Optional[TextChunker],
        cleaner: Optional[TextCleaner],
        table_processor: Optional[TableProcessor],
    ):
        self.chunker = chunker if chunker else TextChunker()
        self.cleaner = cleaner if cleaner else TextCleaner()
        self.table_processor = table_processor if table_processor else TableProcessor()

    def parse(self, file_path: str) -> ParseResult:
        """
        通过文件路径解析文档
        Args:
            file_path: 文档文件路径
        Returns:
            ParseResult: 解析结果
        """
        # 创建docx对象
        doc = docx.Document(file_path)

        # 提取数据
        paragraphs = [
            paragraph.text
            for paragraph in doc.paragraphs
            if paragraph.text.strip() != ""
        ]
        raw_text = "\n\n".join(paragraphs)

        # 提取表格
        tables = self._extract_tables(doc)

        # 处理数据流程
        cleaned_text = self.cleaner.clean(raw_text)
        merged_text = self.table_processor.merge_table(tables)
        chunks = self.chunker.split(cleaned_text, merged_text)

        return ParseResult(
            filename=Path(file_path).name,
            total_pages=1,
            total_chunks=len(chunks),  # 此处的 chunk 数可另作处理，此处保留原意
            table_count=len(merged_text),
            chunks=chunks,
            raw_text=raw_text,
            tables=merged_text,
            metadata={"parser": "docling", "file_path": Path(file_path).name},
        )

    def parse_from_bytes(self, file_bytes: bytes, filename: str) -> ParseResult:
        """
        从字节流解析文档
        Args:
            file_bytes: 文件字节流
            filename: 文件名
        Returns:
            ParseResult: 解析结果
        """
        from io import BytesIO

        doc = docx.Document(BytesIO(file_bytes))

        # 提取数据
        paragraphs = [
            paragraph.text
            for paragraph in doc.paragraphs
            if paragraph.text.strip() != ""
        ]
        raw_text = "\n\n".join(paragraphs)

        # 提取表格
        tables = self._extract_tables(doc)

        # 处理数据流程
        cleaned_text = self.cleaner.clean(raw_text)
        merged_text = self.table_processor.merge_table(tables)
        chunks = self.chunker.split(cleaned_text, merged_text)

        return ParseResult(
            filename=filename,
            total_pages=1,
            total_chunks=len(chunks),  # 此处的 chunk 数可另作处理，此处保留原意
            table_count=len(merged_text),
            chunks=chunks,
            raw_text=raw_text,
            tables=merged_text,
            metadata={"parser": "docling", "file_path": filename},
        )

    def supported_extensions(self) -> list[str]:
        """
        返回支持的文件扩展名
        Returns:
            支持的扩展名列表，如 ['.docx', '.doc']
        """
        return [".docx", ".doc"]

    def _extract_tables(self, doc):
        """提取表格（修正方法名拼写：extract 而非 exstract）"""
        tables = []
        for table_index, table in enumerate(doc.tables):
            # 转换markdown
            md_rows = []
            for row_index, row in enumerate(table.rows):
                # 获取列
                cells = [cell.text.strip() for cell in row.cells]
                md_rows.append("|" + "|".join(cells) + "|")
                # 增加表头
                if row_index == 0:
                    md_rows.append("|" + "|".join(["---"] * len(cells)) + "|")
            tables.append(
                {
                    "page": 1,
                    "rows": len(table.rows),
                    "cols": len(table.columns),
                    "markdown": "\n".join(md_rows),
                }
            )
        return tables
